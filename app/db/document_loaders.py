"""
Custom document loaders for different file types.
"""

import logging
import os

from langchain_community.document_loaders import DirectoryLoader
from langchain_community.document_loaders.base import BaseLoader
from langchain.schema.document import Document

logger = logging.getLogger(__name__)

class DocxLoader(BaseLoader):
    """Loader for .docx files."""
    
    def __init__(self, file_path: str):
        """
        Initialize the DocxLoader.
        
        Args:
            file_path: Path to the .docx file
        """
        self.file_path = file_path
        
    def load(self) -> list[Document]:
        """
        Load the document content.
        
        Returns:
            list containing a single Document with the content of the file
        """
        try:
            # Try using python-docx first
            try:
                import docx
                doc = docx.Document(self.file_path)
                text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                logger.warning(f"Error using python-docx, falling back to docx2txt: {str(e)}")
                # Fall back to docx2txt
                import docx2txt
                text = docx2txt.process(self.file_path)
            
            metadata = {"source": self.file_path}
            return [Document(page_content=text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading .docx file {self.file_path}: {str(e)}")
            raise

class SafeTextLoader(BaseLoader):
    """
    A safer text loader that handles various encoding issues and edge cases.
    """
    
    def __init__(self, file_path: str, encoding: str = 'utf-8'):
        """
        Initialize the SafeTextLoader.
        
        Args:
            file_path: Path to the text file
            encoding: File encoding (default: utf-8)
        """
        self.file_path = file_path
        self.encoding = encoding
        
    def load(self) -> list[Document]:
        """
        Load the text file with robust error handling.
        
        Returns:
            list containing a single Document with the content of the file
        """
        try:
            # Try reading with specified encoding
            with open(self.file_path, 'r', encoding=self.encoding, errors='replace') as f:
                text = f.read()
            
            metadata = {"source": self.file_path}
            return [Document(page_content=text, metadata=metadata)]
            
        except UnicodeDecodeError:
            # Try different encodings if the specified one fails
            try:
                with open(self.file_path, 'r', encoding='latin-1', errors='replace') as f:
                    text = f.read()
                    
                logger.warning(f"Fallback to latin-1 encoding for {self.file_path}")
                metadata = {"source": self.file_path}
                return [Document(page_content=text, metadata=metadata)]
                
            except Exception as backup_error:
                logger.error(f"Error with fallback encoding for {self.file_path}: {str(backup_error)}")
                # Last resort: read as binary and decode with replacement
                with open(self.file_path, 'rb') as f:
                    binary_data = f.read()
                    text = binary_data.decode('utf-8', errors='replace')
                
                metadata = {"source": self.file_path}
                return [Document(page_content=text, metadata=metadata)]
                
        except Exception as e:
            logger.error(f"Error loading text file {self.file_path}: {str(e)}")
            # Instead of raising the exception, return an empty document with error info
            metadata = {
                "source": self.file_path,
                "error": str(e),
                "status": "failed"
            }
            return [Document(page_content=f"Error loading: {str(e)}", metadata=metadata)]

def get_document_loader(file_path: str) -> BaseLoader:
    """
    Get the appropriate document loader based on file extension.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Document loader instance
    """
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    if file_extension == '.docx':
        return DocxLoader(file_path)
    elif file_extension == '.txt':
        return SafeTextLoader(file_path)  # Use our safer loader
    elif file_extension == '.md':
        return SafeTextLoader(file_path)  # Use the safer loader for markdown too
    else:
        # Default to our SafeTextLoader for other extensions
        logger.warning(f"No specific loader for {file_extension}, using SafeTextLoader")
        return SafeTextLoader(file_path)

class CustomDirectoryLoader(DirectoryLoader):
    """Custom directory loader that uses the appropriate loader for each file type."""
    
    def __init__(
        self,
        path: str,
        glob: str = "**/*",
        silent_errors: bool = True,  # Changed to True for better robustness
        load_hidden: bool = False,
    ):
        """Initialize with path."""
        super().__init__(
            path=path,
            glob=glob,
            loader_cls=None,  # We'll override the loading mechanism
            silent_errors=silent_errors,
            load_hidden=load_hidden,
        )
    
    def load(self) -> list[Document]:
        """
        Load documents from the directory, using the appropriate loader for each file.
        
        Returns:
            list of loaded documents
        """
        docs = []
        error_count = 0
        for file_path in self._get_filtered_paths():
            try:
                loader = get_document_loader(file_path)
                loaded_docs = loader.load()
                
                # Check if there was an error in loaded documents (for SafeTextLoader)
                if loaded_docs and "status" in loaded_docs[0].metadata and loaded_docs[0].metadata["status"] == "failed":
                    error_count += 1
                    logger.warning(f"Error loading {file_path}: {loaded_docs[0].page_content}")
                    if not self.silent_errors:
                        raise ValueError(f"Error loading {file_path}: {loaded_docs[0].page_content}")
                    continue
                
                docs.extend(loaded_docs)
                
            except Exception as e:
                error_count += 1
                logger.warning(f"Error loading {file_path}: {str(e)}")
                if not self.silent_errors:
                    raise
        
        if error_count > 0:
            logger.warning(f"Encountered errors in {error_count} files during directory loading")
            
        return docs